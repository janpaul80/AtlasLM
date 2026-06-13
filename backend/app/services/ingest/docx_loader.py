# backend/app/services/ingest/docx_loader.py
"""DOCX loader -> paragraphs + table rows. Requires python-docx."""
from __future__ import annotations
from typing import List
from docx import Document
from .base import ExtractedBlock, block


def load_docx(path: str) -> List[ExtractedBlock]:
    doc = Document(path)
    blocks: List[ExtractedBlock] = []
    offset = 0
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            blocks.append(block(t, char_offset=offset))
            offset += len(t)
    for ti, table in enumerate(doc.tables):
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(block(" | ".join(cells), char_offset=offset))
                offset += sum(len(c) for c in cells)
    return blocks
